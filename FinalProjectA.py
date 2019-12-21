"""Mohammad Zabaneh - Programming Logic and Design - ITEC 1150 - This my final project, and in this project
I will create a Taco recipe cookbook in Word using API including a taco image"""

"""In this section I will import some important libraries that will give me the ability to work with images,
create a word document, and to get the taco dictionary from it's URL"""

# Importing libraries down here
import requests  # imported requests
import docx  # imported python-docs
from PIL import Image, ImageDraw, ImageFont  # imported these libraries to be able wo work on the image
# installed and imported the pillow library

"""In the next section I am going to work on the the image that 
I've downloaded using the pillow library"""

# Working on the image down here
# iimage is the new edited image
image = Image.open('tacorecipe.jpg')  # the pic that I am gonna use (Taco image)
image.thumbnail((700, 700))  # Because the image is big, I resized it to 800/800
iimage = ImageDraw.Draw(image)  # I've used the Imagedraw library to be able to put text on the image
font = ImageFont.truetype('DejaVuSans.ttf', 25)  # using a database for the fonts
iimage.text([50, 400], 'Random Taco Cookbook', fill='black', font=font)  # drawing text on the image
# image.show()  # shows the new edited version of the image
image.save('usingInWord.jpeg')  # This statement saves the image so we can be able to use later in the project

"""In the next section I am going to work on getting the information from it's url
using the "requests" library, and I will create a list to store 3 taco dictionaries in using a loop"""

# working on the dictionaries down here
taco3 = []  # a list to store the dictionaries in

for i in range(3):  # I've used a for loop here so I can get the result from url 3 times (the dictionaries)
    try: # I've added the try and except part just if there was no internet
        taco_recipe_url_data = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()   # using the requests library to open the url link and get the dictionary that we are gonna use
        taco3.append(taco_recipe_url_data)  # adds the results to the list created earlier
    except:
        print('You are not online')  # prints if there is no internet

"""In the next section I am going to work on creating a 
word document using the docx library"""

# Working on the word document down here
word_document = docx.Document()  # This will open a new word document

word_document.add_paragraph('Random Taco Cookbook', 'Title')  # This will add a heading in the new word document

word_document.add_picture('usingInWord.jpeg')  # This statement is gonna use the saved image, and add it to theNewDocument

word_document.add_paragraph('Tai\'s Captures')  # This is gonna add a new paragraph with a name in it

word_document.add_paragraph('https://taco-1150.herokuapp.com/random/?full_taco=true')  # This is gonna add a new paragraph with the taco url in it

word_document.add_paragraph('Mohammad Zabaneh')  # This is gonna add a new paragraph with my name in it

word_document.add_page_break()  # This line of code is gonna add a page break on the first page

for i in range(len(taco3)): # This for loop is gonna pull sections from the taco3 list (the API dictionary)
    word_document.add_paragraph(taco3[0]["seasoning"]["name"], 'heading 1')  #  This line of code gonna grab the name in the seasoning dict
    word_document.add_paragraph(taco3[0]["seasoning"]["recipe"])  # --> first dictionary (seasoning) --> recipe key --> = recipe value
    word_document.add_paragraph(taco3[0]["condiment"]["name"], 'heading 1')  # This line of code gonna grab the name in the condiment dict
    word_document.add_paragraph(taco3[0]["condiment"]["recipe"]) # --> second dictionary (condiment) --> recipe key --> = recipe value
    word_document.add_paragraph(taco3[0]["mixin"]["name"], 'heading 1')  #  This line of code gonna grab the name in the mixin dict
    word_document.add_paragraph(taco3[0]["mixin"]["recipe"]) # --> third dictionary (mixin) --> recipe key --> = recipe value
    word_document.add_paragraph(taco3[0]["base_layer"]["name"], 'heading 1')  # This line of code gonna grab the name in the base layer dict
    word_document.add_paragraph(taco3[0]["base_layer"]["recipe"]) # --> fourth dictionary (base_layer) --> recipe key --> = recipe value
    word_document.add_paragraph(taco3[0]["shell"]["name"], 'heading 1')  # This line of code gonna grab the name in the shell dict
    word_document.add_paragraph(taco3[0]["shell"]["recipe"]) # -->  fifth dictionary (shell) --> recipe key --> = recipe value
    word_document.add_page_break()  # This line of code is gonna add a page break after each recipe
    word_document.save('ThisIsSoCool.docx')  # This line of code is gonna save everything we have added to a word document